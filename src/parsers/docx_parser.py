from pathlib import Path
from docx import Document


def extract_docx(file_path: str) -> dict:
    """
    Extract text and tables from a DOCX document.
    Returns:
        dict: {"text": str, "tables": [{"columns": List[str], "rows": List[List[str]]}]}
    """
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    document = Document(path)

    paragraphs = []
    for paragraph in document.paragraphs:
        text = paragraph.text.strip()
        if text:
            paragraphs.append(text)

    tables_data = []
    for table_idx, table in enumerate(document.tables):
        rows_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)
        
        if rows_data:
            columns = rows_data[0]
            rows = rows_data[1:]
            tables_data.append({
                "columns": columns,
                "rows": rows
            })
            
            paragraphs.append(f"\n--- TABLE {table_idx + 1} ---")
            for r in rows_data:
                paragraphs.append(" | ".join([cell for cell in r if cell.strip()]))

    return {
        "text": "\n".join(paragraphs),
        "tables": tables_data
    }


def extract_docx_text(file_path: str) -> str:
    """
    Extract text from a Word document.
    """
    return extract_docx(file_path)["text"]